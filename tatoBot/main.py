from docx import Document
import os
from deep_translator import GoogleTranslator
import time

def translate_text(text):
    """
    Translate text from Swedish to English using Google Translate via deep-translator.
    """
    if not text.strip():
        return ""
    
    try:
        # Translate from Swedish to English
        translator = GoogleTranslator(source='sv', target='en')
        result = translator.translate(text)
        return result
    except Exception as e:
        print(f"Translation error: {e}")
        # If translation fails, return original text
        return text

def translate_docx(input_file, output_file):
    doc = Document(input_file)
    new_doc = Document()

    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    current_paragraph = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Only translate non-empty paragraphs
            current_paragraph += 1
            print(f"🔄 Translating paragraph {current_paragraph}/{total_paragraphs}: {text[:50]}...")
            
            translated = translate_text(text)
            new_doc.add_paragraph(translated)
            
            # Add a small delay to avoid hitting rate limits
            time.sleep(0.5)
        else:
            new_doc.add_paragraph("")  # Keep empty paragraphs

    new_doc.save(output_file)
    print(f"✅ Saved translated file to: {output_file}")

if __name__ == "__main__":
    input_path = "swedish.docx"  # Put your .docx here
    output_path = "english.docx"

    if not os.path.exists(input_path):
        print(f"❌ File not found: {input_path}")
        print("Please place a file named 'swedish.docx' in the same directory as this script.")
    else:
        print("🔄 Starting translation...")
        print("Using Google Translate API for Swedish to English translation...")
        translate_docx(input_path, output_path)
        print("🎉 Translation completed!")
